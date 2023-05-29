import docx
import os

os.chdir('/home/stuart/Documents/automate_online_materials')
doc = docx.Document('demo.docx')

doc.paragraphs  # list of paragraph objects
print(doc.paragraphs[0].text)  # first paragraph``
p = doc.paragraphs[1]  # second paragraph
print(p.runs)  # list of run objects
print(p.runs[0].text)  # first run
print(p.runs[1].text)  # second run
print(p.runs[2].text)  # third run
print(p.runs[3].text)  # fourth run
print(p.runs[1].bold)  # True
print(p.runs[3].italic)  # True
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'
doc.save('demo2.docx')
print(p.style)  # 'Normal'
p.style = 'Title'
doc.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')  # add_paragraph() method
d.add_paragraph('This is another paragraph.')
d.save('demo4.docx')
p = d.paragraphs[0]  # first paragraph
p.add_run('This is a new run.')  # add_run() method
p.runs
p.runs[1].bold = True
d.save('demo5.docx')


def getText(filename):  # getText() function
    doc = docx.Document(filename)  # create docx object
    fullText = []  # create empty list
    for para in doc.paragraphs:  # loop through paragraphs
        fullText.append(para.text)  # append text to list
        return '\n'.join(fullText)  # join list items with newline char


print(getText('demo.docx'))  # call getText() function
